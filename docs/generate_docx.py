#!/usr/bin/env python3
"""
Generate a .docx file from academic-text-pt.md with FEAU/UNIVAP formatting.
Times New Roman 12pt, A4, margins 3cm top/left, 2cm bottom/right, 1.5 line spacing.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE = os.path.join(SCRIPT_DIR, 'academic-text-pt.md')
ASSETS_DIR = os.path.join(SCRIPT_DIR, 'assets')
OUTPUT_FILE = os.path.join(SCRIPT_DIR, 'TCC_Heloan_Moraes_2026.docx')

FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(12)
LINE_SPACING = 1.5


def setup_styles(doc):
    """Configure document styles for FEAU/UNIVAP formatting."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)

    # Heading 1: bold, uppercase, 12pt, left, no indent
    for level in [1, 2, 3]:
        h = doc.styles[f'Heading {level}']
        h.font.name = FONT_NAME
        h.font.size = FONT_SIZE
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        h.paragraph_format.line_spacing = LINE_SPACING
        h.paragraph_format.space_before = Pt(24 if level == 1 else 18 if level == 2 else 12)
        h.paragraph_format.space_after = Pt(12 if level <= 2 else 8)
        h.paragraph_format.first_line_indent = Cm(0) if level == 1 else Cm(1.25)
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if level == 1:
            h.font.bold = True
            h.font.italic = False
            h.font.underline = False
        elif level == 2:
            h.font.bold = False
            h.font.italic = False
            h.font.underline = True
        else:
            h.font.bold = False
            h.font.italic = True
            h.font.underline = False


def setup_page(doc):
    """Set A4 page size with FEAU margins."""
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)


def parse_markdown(filepath):
    """Read and split the markdown into named sections."""
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()

    sections = {}
    # Extract named preamble sections
    patterns = [
        ('titulo', r'^TITULO\s*\n(.+?)(?=\nRESUMO\b)', re.MULTILINE | re.DOTALL),
        ('resumo', r'^RESUMO\s*\n(.+?)(?=\nPalavras-chave:)', re.MULTILINE | re.DOTALL),
        ('palavras_chave', r'^(Palavras-chave:.+?)(?=\nABSTRACT\b)', re.MULTILINE | re.DOTALL),
        ('abstract', r'^ABSTRACT\s*\n(.+?)(?=\nKeywords:)', re.MULTILINE | re.DOTALL),
        ('keywords', r'^(Keywords:.+?)(?=\n\d+\.\s+[A-Z])', re.MULTILINE | re.DOTALL),
    ]
    for name, pattern, flags in patterns:
        m = re.search(pattern, text, flags)
        if m:
            sections[name] = m.group(1).strip()

    # Extract body (from first numbered chapter to REFERÊNCIAS)
    m = re.search(r'^(1\.\s+INTRODUÇÃO.+?)(?=^REFERÊNCIAS\b)', text, re.MULTILINE | re.DOTALL)
    if m:
        sections['body'] = m.group(1).strip()

    # References
    m = re.search(r'^REFERÊNCIAS\s*\n(.+?)(?=^APÊNDICES\b|^ANEXOS\b|\Z)', text, re.MULTILINE | re.DOTALL)
    if m:
        sections['referencias'] = m.group(1).strip()

    # Appendices
    m = re.search(r'^APÊNDICES\s*\n(.+?)(?=^ANEXOS\b|\Z)', text, re.MULTILINE | re.DOTALL)
    if m:
        sections['apendices'] = m.group(1).strip()

    # Annexes
    m = re.search(r'^ANEXOS\s*\n(.+?)$', text, re.MULTILINE | re.DOTALL)
    if m:
        sections['anexos'] = m.group(1).strip()

    return sections


def add_inline_formatting(paragraph, text):
    """Add a text run with inline markdown formatting (bold, italic, code)."""
    # Process inline patterns: **bold**, *italic*, `code`
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def add_paragraph_text(doc, text, style='Normal', alignment=None, indent=True,
                       bold=False, italic=False, font_size=None, space_after=None,
                       line_spacing=None):
    """Add a paragraph with inline formatting."""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.paragraph_format.alignment = alignment
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing

    add_inline_formatting(p, text)

    if bold or italic or font_size:
        for run in p.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if font_size:
                run.font.size = font_size
    return p


def add_image(doc, img_path, caption=None, width=None):
    """Add an image with optional caption."""
    if not os.path.exists(img_path):
        add_paragraph_text(doc, f'[Imagem não encontrada: {img_path}]',
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        return

    if caption:
        add_paragraph_text(doc, caption, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           indent=False, bold=True, font_size=Pt(10), space_after=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    w = width or Cm(14)
    run.add_picture(img_path, width=w)

    add_paragraph_text(doc, 'Fonte: O autor.', alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       indent=False, italic=True, font_size=Pt(10), space_after=12)


def process_body(doc, body_text):
    """Process the body markdown into Word paragraphs."""
    lines = body_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ''

    while i < len(lines):
        line = lines[i]

        # Code block start/end
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lang = line.strip()[3:].strip()
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Image: ![caption](path)
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            # Resolve relative path
            if not os.path.isabs(img_path):
                img_path = os.path.join(SCRIPT_DIR, img_path)
            add_image(doc, img_path, caption=caption if caption else None)
            i += 1
            # Check for source line after image
            if i < len(lines) and lines[i].strip().startswith('*') and ('Fonte' in lines[i] or 'fonte' in lines[i]):
                src_text = lines[i].strip().strip('*').strip()
                add_paragraph_text(doc, src_text, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                   indent=False, italic=True, font_size=Pt(10), space_after=12)
                i += 1
            continue

        # H1: N. TITLE (all caps)
        h1_match = re.match(r'^(\d+)\.\s+([A-ZÀÁÂÃÉÊÍÓÔÕÚ\s]+)$', line.strip())
        if h1_match:
            title = line.strip()
            doc.add_heading(title, level=1)
            i += 1
            continue

        # H2: N.N Title
        h2_match = re.match(r'^(\d+\.\d+)\s+(.+)$', line.strip())
        if h2_match:
            title = line.strip()
            doc.add_heading(title, level=2)
            i += 1
            continue

        # H3: N.N.N Title
        h3_match = re.match(r'^(\d+\.\d+\.\d+)\s+(.+)$', line.strip())
        if h3_match:
            title = line.strip()
            doc.add_heading(title, level=3)
            i += 1
            continue

        # Table: lines starting with |
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, table_lines)
            continue

        # Unordered list: - item
        if re.match(r'^[-*]\s+', line.strip()):
            items = []
            while i < len(lines) and re.match(r'^[-*]\s+', lines[i].strip()):
                items.append(re.sub(r'^[-*]\s+', '', lines[i].strip()))
                i += 1
            for item in items:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.first_line_indent = Cm(0)
                add_inline_formatting(p, item)
            continue

        # Ordered list: (i), (ii), etc or 1. item
        if re.match(r'^\([ivxlcdm]+\)\s+', line.strip()):
            items = []
            while i < len(lines) and re.match(r'^\([ivxlcdm]+\)\s+', lines[i].strip()):
                items.append(lines[i].strip())
                i += 1
            for item in items:
                add_paragraph_text(doc, item, indent=True)
            continue

        # Figure source italic line
        if line.strip().startswith('*') and ('Fonte' in line or 'fonte' in line or 'Figura' in line or 'Tabela' in line):
            text = line.strip().strip('*').strip()
            add_paragraph_text(doc, text, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               indent=False, italic=True, font_size=Pt(10), space_after=12)
            i += 1
            continue

        # Regular paragraph - collect continuation lines
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            if not next_line.strip():
                break
            if next_line.strip().startswith('```'):
                break
            if re.match(r'^!\[', next_line.strip()):
                break
            if re.match(r'^\d+\.\s+[A-ZÀÁÂÃÉÊÍÓÔÕÚ]', next_line.strip()):
                break
            if re.match(r'^\d+\.\d+', next_line.strip()):
                break
            if next_line.strip().startswith('|'):
                break
            if re.match(r'^[-*]\s+', next_line.strip()):
                break
            if re.match(r'^\([ivxlcdm]+\)', next_line.strip()):
                break
            para_lines.append(next_line)
            i += 1

        full_text = ' '.join(l.strip() for l in para_lines)
        add_paragraph_text(doc, full_text)


def add_table(doc, table_lines):
    """Parse markdown table lines and add a Word table."""
    # Filter out separator lines (|---|---|)
    data_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:]+\|$', l)]
    if not data_lines:
        return

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                add_inline_formatting(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = FONT_NAME
                # Bold for header row
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True

    # Set table width
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(15 / num_cols)


def add_cover_page(doc, sections):
    """Add the cover page."""
    titulo = sections.get('titulo', 'Título do Trabalho')

    # University name
    for _ in range(3):
        doc.add_paragraph()

    add_paragraph_text(doc, 'UNIVERSIDADE DO VALE DO PARAÍBA',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                       bold=True, font_size=Pt(14))
    add_paragraph_text(doc, 'FACULDADE DE ENGENHARIAS, ARQUITETURA E URBANISMO',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                       bold=True, font_size=Pt(13), space_after=6)
    add_paragraph_text(doc, 'Engenharia da Computação',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                       font_size=Pt(12))

    for _ in range(5):
        doc.add_paragraph()

    add_paragraph_text(doc, titulo.upper(),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                       bold=True, font_size=Pt(14))

    for _ in range(5):
        doc.add_paragraph()

    add_paragraph_text(doc, 'Heloan Barbosa de Moraes',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    for _ in range(2):
        doc.add_paragraph()

    add_paragraph_text(doc, 'Orientador: Prof. Me. Evânio Ramos Nicoleit',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    for _ in range(5):
        doc.add_paragraph()

    add_paragraph_text(doc, 'São José dos Campos — SP',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_paragraph_text(doc, '2026',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    doc.add_page_break()


def add_title_page(doc, sections):
    """Add the title page (folha de rosto)."""
    titulo = sections.get('titulo', 'Título do Trabalho')

    for _ in range(3):
        doc.add_paragraph()

    add_paragraph_text(doc, 'UNIVERSIDADE DO VALE DO PARAÍBA',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                       bold=True, font_size=Pt(14))
    add_paragraph_text(doc, 'FACULDADE DE ENGENHARIAS, ARQUITETURA E URBANISMO',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                       bold=True, font_size=Pt(13), space_after=6)
    add_paragraph_text(doc, 'Engenharia da Computação',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    for _ in range(3):
        doc.add_paragraph()

    add_paragraph_text(doc, titulo.upper(),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False,
                       bold=True, font_size=Pt(14))

    add_paragraph_text(doc, 'Heloan Barbosa de Moraes',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    for _ in range(3):
        doc.add_paragraph()

    # Natureza do trabalho (right-aligned block)
    natureza = ('Trabalho de Conclusão de Curso apresentado ao Programa de '
                'Graduação em Engenharia da Computação da Universidade do '
                'Vale do Paraíba, como complemento às exigências para '
                'obtenção do título de Bacharel em Engenharia da Computação.')
    p = add_paragraph_text(doc, natureza, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           indent=False, font_size=Pt(11), line_spacing=1.0)
    p.paragraph_format.left_indent = Cm(8)

    doc.add_paragraph()
    p = add_paragraph_text(doc, 'Orientador: Prof. Me. Evânio Ramos Nicoleit',
                           indent=False, font_size=Pt(11))
    p.paragraph_format.left_indent = Cm(8)

    for _ in range(6):
        doc.add_paragraph()

    add_paragraph_text(doc, 'São José dos Campos — SP',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_paragraph_text(doc, '2026',
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    doc.add_page_break()


def add_resumo_page(doc, sections):
    """Add RESUMO and ABSTRACT pages."""
    # RESUMO
    p = doc.add_heading('RESUMO', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)

    resumo = sections.get('resumo', '')
    p = add_paragraph_text(doc, resumo, indent=False, line_spacing=1.0)

    palavras = sections.get('palavras_chave', '')
    if palavras:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        # Split on first colon
        parts = palavras.split(':', 1)
        run = p.add_run(parts[0] + ':')
        run.bold = True
        if len(parts) > 1:
            add_inline_formatting(p, parts[1])

    doc.add_page_break()

    # ABSTRACT
    p = doc.add_heading('ABSTRACT', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)

    abstract = sections.get('abstract', '')
    p = add_paragraph_text(doc, abstract, indent=False, line_spacing=1.0)

    keywords = sections.get('keywords', '')
    if keywords:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        parts = keywords.split(':', 1)
        run = p.add_run(parts[0] + ':')
        run.bold = True
        if len(parts) > 1:
            add_inline_formatting(p, parts[1])

    doc.add_page_break()


def add_references(doc, ref_text):
    """Add REFERÊNCIAS section."""
    p = doc.add_heading('REFERÊNCIAS', level=1)
    p.paragraph_format.space_before = Pt(0)

    for line in ref_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        p = add_paragraph_text(doc, line, indent=False, line_spacing=1.0,
                               space_after=6)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_appendices(doc, text, title):
    """Add APÊNDICES or ANEXOS section."""
    p = doc.add_heading(title, level=1)
    p.paragraph_format.space_before = Pt(0)
    process_body(doc, text)


def main():
    print(f'Reading {MD_FILE}...')
    sections = parse_markdown(MD_FILE)

    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    print('Adding cover page...')
    add_cover_page(doc, sections)

    print('Adding title page...')
    add_title_page(doc, sections)

    print('Adding RESUMO / ABSTRACT...')
    add_resumo_page(doc, sections)

    print('Processing body...')
    if 'body' in sections:
        process_body(doc, sections['body'])

    if 'referencias' in sections:
        doc.add_page_break()
        print('Adding REFERÊNCIAS...')
        add_references(doc, sections['referencias'])

    if 'apendices' in sections:
        doc.add_page_break()
        print('Adding APÊNDICES...')
        add_appendices(doc, sections['apendices'], 'APÊNDICES')

    if 'anexos' in sections:
        doc.add_page_break()
        print('Adding ANEXOS...')
        add_appendices(doc, sections['anexos'], 'ANEXOS')

    print(f'Saving {OUTPUT_FILE}...')
    doc.save(OUTPUT_FILE)
    print(f'Done! File saved: {OUTPUT_FILE}')


if __name__ == '__main__':
    main()
